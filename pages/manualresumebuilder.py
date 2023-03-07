import streamlit as st
st.set_page_config(page_title="19th Street | Manual Resume Builder", page_icon="📜", layout="wide", initial_sidebar_state='collapsed')
import streamlit as st
import openai
from docx import Document
from PIL import Image
from streamlit_extras.switch_page_button import switch_page
import pdfkit
from htmldocx import HtmlToDocx
import iomu
from docx.shared import Inches, Cm
import streamlit as st
import os
import openai
from streamlit_extras.switch_page_button import switch_page
import PyPDF2
import datetime
import requests
from streamlit_option_menu import option_menu
import extra_streamlit_components as stx
from st_btn_select import st_btn_select
import pyrebase




firebaseconfig = {
    "apiKey": "AIzaSyDCHY-GB5WCd0V6o4psrasOYZL_F7xcODM",
    "authDomain": "nineteenth-street.firebaseapp.com",
    "projectId": "nineteenth-street",
    "storageBucket": "nineteenth-street.appspot.com",
    "messagingSenderId": "964724806859",
    "appId": "1:964724806859:web:010841fc337f30b50cb74e",
    "measurementId": "G-N3TMC7M1WT",
    "databaseURL": "https://nineteenth-street-default-rtdb.firebaseio.com"
}

if __name__ == "__main__":
    def main(user: object):
        hide_menu_style = """
                 <style>
                 #MainMenu {visibility: hidden;}
                 .css-j7qwjs {visibility: hidden;}
                 footer {visibility: hidden;}
                 </style>
                 """
        st.markdown(hide_menu_style, unsafe_allow_html=True)
        hide_streamlit_style = """
                      <style>
                      div[class='css-4z1n4l ehezqtx5']{
                        background: rgba(0, 0, 0, 0.3);
                        color: #fff;
                        border-radius: 10px;
                        backdrop-filter: blur(10px);
                        height: 40px;
                        max-width: 200px;
                        position: fixed;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%);
                        width: 50%;
                      }
                      </style>
                      """
        st.markdown(hide_streamlit_style, unsafe_allow_html=True)

        st.markdown(f"<h2 style='text-align: center; font-family: Sans-Serif; color:black>Coming Sooner</h2>",
                    unsafe_allow_html=True)

        hide_img_fs = '''
        <style>
        button[title="View fullscreen"]{
            visibility: hidden;}


        </style>
        '''

        st.markdown(hide_img_fs, unsafe_allow_html=True)

        st.markdown("""
            <style>

            css-17z41qg {
                background-color: #eeeeee;
            }
            </style>
            """, unsafe_allow_html=True)

        st.markdown(
            f"<center> <h1 style='font-family: Sans-Serif; font-weight:lighter; color: white'><span style='background: -webkit-gradient(linear,left top,right bottom,from(#34C800), to(#FE0000));-webkit-background-clip:text;-webkit-text-fill-color: transparent;'>19th street</span> Manual Resume Builder.</h1>",
            unsafe_allow_html=True)
        st.markdown(
            f"<center> <h4 style='font-family: Sans-Serif; font-weight:lighter; color: white'>Enter the details. Let AI do the rest.</h4>",
            unsafe_allow_html=True)
        st.subheader("")
        st.subheader("")

        tab1, tab2, tab3, tab4, tab5 = st.tabs(
            ["\u2001Basics\u2001", "\u2001\u2001Experience\u2001\u2001", "\u2001\u2001Projects\u2001\u2001",
             "\u2001\u2001Skills\u2001\u2001", "\u2001\u2001Result\u2001\u2001"])

        with tab1:
            st.header("Details")
            col1, col2, col3 = st.columns([1, 1, 1])
            with col1:
                CandidateName = st.text_input(
                    'Name',
                    placeholder='Name ',
                    key='Name'
                )

                CandidatePhone = st.text_input(
                    'Phone',
                    placeholder='Phone Number',
                    key='Phone'
                )

                CandidateEmail = st.text_input(
                    'Email',
                    placeholder='Email ',
                    key='Email'
                )
            with col2:
                st.write("")
            with col3:
                st.write("")

        with tab2:
            st.header("Experiences")
            st.markdown(
                f"<h6 style='font-family: Sans-Serif; font-weight:lighter; color: white'>Don't worry about being perfect. Just write briefly about what you did and we'll clean it up and generate bullet points out of it.</h6>",
                unsafe_allow_html=True)

            col1, col2 = st.columns([0.75, 1])
            with col1:
                st.markdown("""
                    <style>

                    div[inputmode="text"]{
                    border-radius:5px;
                    -moz-border-radius:5px;
                    -webkit-border-radius:5px;
                    border: 0 none;
                    outline: none;
                    }

                    .st-br{
                    border-radius:10px;
                    -moz-border-radius:10px;
                    -webkit-border-radius:10px;

                    outline: none;
                    }

                    ul.streamlit-expander {
                    border-radius: 12.5px !important;
                    }

                    .stTextInput [data-baseweb=base-input] {
                        background-color: #0d0d0d;
                        -webkit-text-fill-color: white;
                    }

                    .stTextInput [data-baseweb=base-input] [disabled=""]{
                        background-color: #0d0d0d;
                        -webkit-text-fill-color: white;
                    }

                    .stTextArea [data-baseweb=base-input] {
                        background-color: #0d0d0d;
                        -webkit-text-fill-color: white;
                        background-color: #0d0d0d;
                          width: 100%;
                          font-size: 5em;
                          outline: none;
                          position: relative;
                    }

                    .stTextArea [data-baseweb=base-input] [disabled=""]{
                        background-color: #0d0d0d;
                        -webkit-text-fill-color: white;
                         background-color: #0d0d0d;
                          width: 100%;
                          font-size: 5em;
                          outline: none;
                          position: relative;
                    }

                    </style>
                    """, unsafe_allow_html=True)

                with st.expander("Experience 1", expanded=False):
                    Experience1Name = st.text_input(
                        '',
                        placeholder='Position at Company',
                        key='Experience1'
                    )

                    Experience1Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Experience 1 Detail'
                    )
                with st.expander("Experience 2", expanded=False):
                    Experience2Name = st.text_input(
                        '',
                        placeholder='Position at Company',
                        key='Experience2'
                    )
                    Experience2Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Experience 2 Detail'
                    )
                with st.expander("Experience 3", expanded=False):
                    Experience3Name = st.text_input(
                        '',
                        placeholder='Position at Company',
                        key='Experience3'
                    )
                    Experience3Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Experience 3 Detail'
                    )
                with st.expander("Experience 4", expanded=False):
                    Experience4Name = st.text_input(
                        '',
                        placeholder='  Position at Company',
                        key='Experience4'
                    )
                    Experience4Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Experience 4 Detail'
                    )
                with st.expander("Experience 5", expanded=False):
                    Experience5Name = st.text_input(
                        '',
                        placeholder='Position at Company',
                        key='Experience5'
                    )
                    Experience5Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Experience 5 Detail'
                    )

            with col2:
                st.write("")

        with tab3:
            col1tab3, col2tab3 = st.columns([0.75, 1])
            with col1tab3:
                st.header("Projects")
                with st.expander("Project 1", expanded=False):
                    col1Project1, col2Project1 = st.columns([1, 1])
                    with col1Project1:
                        Projec1Name = st.text_input(
                            '',
                            placeholder='Project Name',
                            key='Project1Name'
                        )
                    with col2Project1:
                        Project1Link = st.text_input(
                            '',
                            placeholder='Project Link',
                            key='Project1Link'
                        )

                    Project1Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Project 1 Detail'
                    )
                with st.expander("Project 2", expanded=False):
                    col1Project2, col2Project2 = st.columns([1, 1])
                    with col1Project2:
                        Projec2Name = st.text_input(
                            '',
                            placeholder='Project Name',
                            key='Project2Name'
                        )
                    with col2Project2:
                        Project1Link = st.text_input(
                            '',
                            placeholder='Project Link',
                            key='Project2Link'
                        )

                    Project2Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Project 2 Detail'
                    )
                with st.expander("Project 3", expanded=False):
                    col1Project3, col2Project3 = st.columns([1, 1])
                    with col1Project3:
                        Projec3Name = st.text_input(
                            '',
                            placeholder='Project Name',
                            key='Project3Name'
                        )
                    with col2Project3:
                        Projec3Link = st.text_input(
                            '',
                            placeholder='Project Link',
                            key='Project3Link'
                        )

                    Project3Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Project 3 Detail'
                    )
                with st.expander("Project 4", expanded=False):
                    col1Project4, col2Project4 = st.columns([1, 1])
                    with col1Project4:
                        Projec4Name = st.text_input(
                            '',
                            placeholder='Project Name',
                            key='Project4Name'
                        )
                    with col2Project4:
                        Projec4Link = st.text_input(
                            '',
                            placeholder='Project Link',
                            key='Project4Link'
                        )

                    Project4Description = st.text_area(
                        '',
                        placeholder='Description',
                        key='Project 4 Detail'
                    )

            with col2tab3:
                st.write("")

        with tab4:
            st.header("Skills")
            if 'Skills' not in st.session_state:
                if st.button("Add Skills"):
                    response = openai.Completion.create(
                        model="text-davinci-003",
                        prompt=f"The following is some experience of a job seeker.\n\n{Experience1Name}\n{Experience1Description}\n\n{Experience2Name}\n{Experience2Description}\n\n{Experience3Name}\n{Experience3Description}\n\n{Experience4Name}\n{Experience4Description} \n What kind of technical skills they have? List them as spearated by commas.\n",
                        temperature=0.7,
                        max_tokens=80,
                        top_p=1,
                        frequency_penalty=0,
                        presence_penalty=0
                    )

                    PreSkills = response["choices"][0]["text"].split(",")
                    st.session_state['Skills'] = ','.join(PreSkills)

                    FinalSkills = st.text_input(
                        'Caption goes here',
                        placeholder='Placeholder goes here',
                        help='Help message goes here',
                        value=st.session_state['Skills']
                    )
            else:
                FinalSkills = st.text_input(
                    'Caption goes here',
                    placeholder='Placeholder goes here',
                    help='Help message goes here',
                    value=st.session_state['Skills'])

        with tab5:
            st.subheader("Choose the role/industry you wish to tailor your resume to")

            @st.cache_data
            def GettingJobTitles(Experience1Name, Experience1Description, Experience2Name, Experience2Description,
                                 Experience3Name, Experience3Description, Experience4Name, Experience4Description):
                responseProjects = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system",
                         "content": "You are an AI Assistant that ouputs 7 relevant job titles in addition what the job seeker has already done. Your response is a list of job titles separated by commas. You response doesn't include any extra fluff."},
                        {"role": "user",
                         "content": f"The following is some experience of a job seeker.\n\n{Experience1Name}\n{Experience1Description}\n\n{Experience2Name}\n{Experience2Description}\n\n{Experience3Name}\n{Experience3Description}\n\n{Experience4Name}\n{Experience4Description} \n"}])

                SuggestedJobtitlesResponse = responseProjects["choices"][0]["message"]["content"]
                return SuggestedJobtitlesResponse

            SuggestedJobTitles = GettingJobTitles(Experience1Name, Experience1Description, Experience2Name,
                                                  Experience2Description, Experience3Name, Experience3Description,
                                                  Experience4Name, Experience4Description)
            ChosenJobTitle = st.selectbox(
                '',
                SuggestedJobTitles.split(","),
                key="ChosenJobTitle"
            )

            if st.button("Proceed →"):
                NewExperienceOneDescriptionResponse = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system",
                         "content": f"""
                         You're take in resume experience and rewrite them to sound like an experienced {ChosenJobTitle}.
                         Your response isn't in a paragraph form.
                         You start every bullet point with a semi colon.
                         Your response should resemble this format:
                         ;point 1
                         ;point 2
                         ;point 3
                         """},
                        {"role": "user",
                         "content": f"The following is description of experience of a job seeker.\n{Experience1Description}. Make it sound they're an experience {ChosenJobTitle}"}])
                NewExperienceOneDescription = NewExperienceOneDescriptionResponse["choices"][0]["message"]["content"]

                NewExperienceTwoDescriptionResponse = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system",
                         "content": f"""
                         You're take in resume experience and rewrite them to sound like an experienced {ChosenJobTitle}.
                         Your response isn't in a paragraph form.
                         You start every bullet point with a semi colon.
                         Your response should resemble this format:
                         ;point 1
                         ;point 2
                         ;point 3
                         """},
                        {"role": "user",
                         "content": f"The following is description of experience of a job seeker.\n{Experience2Description}. Make it sound they're an experience {ChosenJobTitle}"}])
                NewExperienceTwoDescription = NewExperienceTwoDescriptionResponse["choices"][0]["message"]["content"]

                NewExperienceThreeDescriptionResponse = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system",
                         "content": f"""
                         You're take in resume experience and rewrite them to sound like an experienced {ChosenJobTitle}.
                         Your response isn't in a paragraph form.
                         You start every bullet point with a semi colon.
                         Your response should resemble this format:
                         ;point 1
                         ;point 2
                         ;point 3
                         """},
                        {"role": "user",
                         "content": f"The following is description of experience of a job seeker.\n{Experience3Description}. Make it sound they're an experience {ChosenJobTitle}"}])
                NewExperienceThreeDescription = NewExperienceThreeDescriptionResponse["choices"][0]["message"][
                    "content"]

                NewExperienceFourDescriptionResponse = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system",
                         "content": f"""
                         You're take in resume experience and rewrite them to sound like an experienced {ChosenJobTitle}.
                         Your response isn't in a paragraph form.
                         You start every bullet point with a semi colon.
                         Your response should resemble this format:
                         ;point 1
                         ;point 2 
                         ;point 3 
                         """},
                        {"role": "user",
                         "content": f"The following is description of experience of a job seeker.\n{Experience4Description}. Make it sound they're an experience {ChosenJobTitle}"}])
                NewExperienceFourDescription = NewExperienceFourDescriptionResponse["choices"][0]["message"]["content"]

                if Projec1Name:
                    NewProjectOneDescriptionResponse = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system",
                             "content": f"""
                                         You're take in resume project and rewrite it as a resume bullet point to sound like an experienced {ChosenJobTitle} in less than 30 words.
                                     """},
                            {"role": "user",
                             "content": f"The following is description of a project of a job seeker.\n{Project1Description}. Make it sound they're an experienced {ChosenJobTitle}"}])
                    NewProjectOneDescription = NewProjectOneDescriptionResponse["choices"][0]["message"]["content"]

                if Projec2Name:
                    NewProjectTwoDescriptionResponse = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system",
                             "content": f"""
                                         You're take in resume project and rewrite it as a resume bullet point to sound like an experienced {ChosenJobTitle} in less than 30 words.
                                         """},
                            {"role": "user",
                             "content": f"The following is description of a project of a job seeker.\n{Project2Description}. Make it sound they're an experienced {ChosenJobTitle}"}])
                    NewProjectTwoDescription = NewProjectTwoDescriptionResponse["choices"][0]["message"]["content"]

                if Projec3Name:
                    NewProjectThreeDescriptionResponse = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system",
                             "content": f"""
                                         You're take in resume project and rewrite it as a resume bullet point to sound like an experienced {ChosenJobTitle} in less than 30 words.
                                         """},
                            {"role": "user",
                             "content": f"The following is description of a project of a job seeker.\n{Project3Description}. Make it sound they're an experienced {ChosenJobTitle}"}])
                    NewProjectThreeDescription = NewProjectThreeDescriptionResponse["choices"][0]["message"]["content"]

                if Projec4Name:
                    NewProjectFourDescriptionResponse = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system",
                             "content": f"""
                                         You're take in resume project and rewrite it as a resume bullet point to sound like an experienced {ChosenJobTitle} in less than 30 words.
                                         """},
                            {"role": "user",
                             "content": f"The following is description of a project of a job seeker.\n{Project4Description}. Make it sound they're an experienced {ChosenJobTitle}"}])
                    NewProjectFourDescription = NewProjectFourDescriptionResponse["choices"][0]["message"]["content"]

                html_string = ""

                html_string += "<h3 style='font-size:15px;align:center'>" + CandidatePhone + ' | ' + CandidateEmail + "</h3>"

                html_string += "<h3>" + Experience1Name + "</h3>"
                for item in NewExperienceOneDescription.split(";")[1:]:
                    html_string += "<li>" + item + "</li>"

                html_string += "<h3>" + Experience2Name + "</h3>"
                for item in NewExperienceTwoDescription.split(";")[1:]:
                    html_string += "<li>" + item + "</li>"

                html_string += "<h3>" + Experience3Name + "</h3>"
                for item in NewExperienceThreeDescription.split(";")[1:]:
                    html_string += "<li>" + item + "</li>"

                html_string += "<h3> Skills:" + FinalSkills + "</h3>"

                if Projec1Name:
                    html_string2 = ""

                    html_string2 += "<h3>" + Projec1Name + "</h3>"
                    html_string2 += "<li>" + NewProjectOneDescription + "</li>"

                    html_string2 += "<h3>" + Projec2Name + "</h3>"
                    html_string2 += "<li>" + NewProjectTwoDescription + "</li>"

                    html_string2 += "<h3>" + Projec3Name + "</h3>"
                    html_string2 += "<li>" + NewProjectThreeDescription + "</li>"

                document = Document()
                heading = document.add_heading(f'{CandidateName}', 0)
                sections = document.sections
                for section in sections:
                    section.top_margin = Cm(0.5)
                    section.bottom_margin = Cm(0.5)
                    section.left_margin = Cm(1.5)
                    section.right_margin = Cm(1.5)

                new_parser = HtmlToDocx()
                headingExperiences = document.add_heading(f'Experiences', 1)
                new_parser.add_html_to_document(html_string, document)
                if Projec1Name:
                    headingProjects = document.add_heading(f'Projects', 1)
                    new_parser.add_html_to_document(html_string2, document)

                # do more stuff to document
                document.save('your_file_name')

                # document.save('your_file_name')

                doc_download = document

                bio = io.BytesIO()
                doc_download.save(bio)
                if doc_download:
                    st.download_button(
                        label="Click here to download",
                        data=bio.getvalue(),
                        file_name="19th_Street_Resume_Edits.docx",
                        mime="docx"
                    )


@st.cache(allow_output_mutation=True)
def get_manager():
    return stx.CookieManager()


cookie_manager = get_manager()


def set_code(code: str):
    st.experimental_set_query_params(code=code)
    cookie = "queryParamCode"
    val = str(code)
    cookie_manager.set(cookie, val, expires_at=datetime.datetime(year=2024, month=2, day=2))


col1form, col2form, col3form = st.columns([0.25, 1, 0.25])
with col1form:
    st.write("")
with col2form:
    def login_form(auth):

        st.subheader("All Cookies:")
        cookies = cookie_manager.get_all()
        st.write(cookies)
        email = st.text_input(
            label="email", placeholder="fullname@gmail.com")
        password = st.text_input(
            label="password", placeholder="password", type="password")

        if st.button("Login"):
            try:
                user = auth.sign_in_with_email_and_password(email, password)
                st.session_state['user'] = user
                st.experimental_rerun()
            except requests.HTTPError as exception:
                st.write(exception)
        if st.button("Forgot Password", key="forgotpassword"):
            auth.send_password_reset_email("email")


    def Signup_form(auth):
        email = st.text_input(
            label="email", placeholder="fullname@gmail.com")
        password = st.text_input(
            label="password", placeholder="password", type="password")

        if st.button("login"):
            try:
                user = auth.sign_in_with_email_and_password(email, password)
                st.session_state['user'] = user
                st.experimental_rerun()
            except requests.HTTPError as exception:
                st.write(exception)
        if st.button("Forgot Password", key="forgotpassword"):
            auth.send_password_reset_email("email")

with col3form:
    st.write("")


def logout():
    del st.session_state['user']
    st.experimental_set_query_params(code="/logout")


def get_user_token(auth, refreshToken: object):
    user = auth.get_account_info(refreshToken['idToken'])

    user = {
        "email": user['users'][0]['email'],
        "refreshToken": refreshToken['refreshToken'],
        "idToken": refreshToken['idToken']
    }

    st.session_state['user'] = user

    return user


def refresh_session_token(auth, code: str):
    try:
        return auth.refresh(code)
    except:
        return "fail to refresh"


firebase = pyrebase.initialize_app(firebaseconfig)

auth = firebase.auth()

# authentification
if "user" not in st.session_state:
    st.session_state['user'] = None

if st.session_state['user'] is None:
    try:
        # code = st.experimental_get_query_params()['code'][0]
        code = cookie_manager.get(cookie="queryParamCode")

        refreshToken = refresh_session_token(auth=auth, code=code)

        if refreshToken == 'fail to refresh':
            raise (ValueError)

        user = get_user_token(auth, refreshToken=refreshToken)

        main(user=user)
    except:
        st.title("Login")
        login_form(auth)

else:
    main(user=st.session_state['user'])